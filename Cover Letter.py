# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""
import pandas as pd
import numpy as np
import docx

df_saved_texts = pd.read_excel('Saved Texts.xlsx', index_col=[0])
company = 'Ontario Teachers Pension Plan'
gmap_location = "8888 University Dr W, Burnaby, BC V5A 1S6"
job_position = 'Private Equity Analyst'
team = "Private Equity"

selection_1 = 'Passion in Capital Markets'
selection_2 = 'GMG - Equity Research'
selection_3 = 'BCI - Public Markets Risk Management'

description_1 = df_saved_texts.loc[selection_1]['Description']
description_2 = df_saved_texts.loc[selection_2]['Description']
description_3 = df_saved_texts.loc[selection_3]['Description']


p2 = df_saved_texts.loc[selection_1]['Text']
p3 = df_saved_texts.loc[selection_2]['Text']
p4 = df_saved_texts.loc[selection_3]['Text']

gmap_address = gmap_location.split(',')[0].lstrip(' ')
gmap_city_prov_code = gmap_location.split(',')[1].lstrip(' ') + gmap_location.split(',')[2]

document = docx.Document(r'Cover Letter Template.docx')
style = document.styles['Normal']
font = style.font
font.name = 'Helvetica'
font.size = docx.shared.Pt(10)

def replace_text(text, text_to_replace):
    for i in document.paragraphs:
        if text in i.text:
            i.text = i.text.replace(text, text_to_replace)

paragraph_intro = 'I hope your week is going well. My name is Hans Suherman, I am an undergraduate student from Simon Fraser University (Graduating December 2022) with '+ description_1 + ', '+ description_2 +', and ' + description_3 + '. With this letter, I intend to express my interest in joining ' + company + ' because I believe that I can add great value to the '+ team +' team.'

replace_text('Company', company)
replace_text('Address', gmap_address)
replace_text('City_Prov_Code', gmap_city_prov_code)

replace_text('Job_Position', job_position + ' Position')
replace_text('P1', paragraph_intro)
replace_text('P2', p2)
replace_text('P3', p3)
replace_text('P4', p4)
#replace_text('P5', p5)
#replace_text('P6', p6)


for i in document.paragraphs:
    print(i.text)

document.save('Test.docx')
